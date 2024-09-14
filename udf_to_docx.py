import sys
import os
import xml.etree.ElementTree as ET
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import base64
import io
import zipfile

def udf_to_docx(udf_file, docx_file):
    # Open the UDF file and read the content.xml file
    with zipfile.ZipFile(udf_file, 'r') as z:
        if 'content.xml' in z.namelist():
            with z.open('content.xml') as content_file:
                tree = ET.parse(content_file)
                root = tree.getroot()
        else:
            print("The 'content.xml' file could not be found in the UDF file.")
            exit()

    # Create a new Word document
    document = Document()

    # Create a dictionary for style definitions
    styles = {}

    # Retrieve style information
    styles_element = root.find('styles')
    if styles_element is not None:
        for style in styles_element.findall('style'):
            style_name = style.get('name')
            style_attributes = {
                'family': style.get('family'),
                'size': int(style.get('size', 12)),
                'bold': style.get('bold', 'false') == 'true',
                'italic': style.get('italic', 'false') == 'true',
                'foreground': int(style.get('foreground', '-13421773')),
            }
            styles[style_name] = style_attributes

    # Retrieve content text
    content_element = root.find('content')
    if content_element is not None:
        content_text = content_element.text
        if content_text.startswith('<![CDATA[') and content_text.endswith(']]>'):
            content_text = content_text[9:-3]
    else:
        print("'content' could not be found in the XML.")
        exit()

    # Process the 'elements' section
    elements_element = root.find('elements')
    if elements_element is not None:
        for elem in elements_element:
            if elem.tag == 'paragraph':
                # Create the paragraph
                paragraph = document.add_paragraph()

                # Set paragraph alignment
                alignment = elem.get('Alignment', '0')
                if alignment == '0':
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                elif alignment == '1':
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif alignment == '2':
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                elif alignment == '3':
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

                # Process the paragraph content
                for child in elem:
                    if child.tag == 'content':
                        # Get and format the text
                        start_offset = int(child.get('startOffset', '0'))
                        length = int(child.get('length', '0'))
                        text = content_text[start_offset:start_offset+length]

                        run = paragraph.add_run(text)

                        # Set the font size
                        size = int(child.get('size', '12'))
                        run.font.size = Pt(size)

                        # Set bold and italic options
                        if child.get('bold', 'false') == 'true':
                            run.bold = True
                        if child.get('italic', 'false') == 'true':
                            run.italic = True

                        # Set the font family
                        font_family = child.get('family')
                        if font_family:
                            run.font.name = font_family

                        # Set the color
                        foreground = child.get('foreground')
                        if foreground:
                            color_int = int(foreground)
                            if color_int < 0:
                                color_int += 2**32  # Convert negative values to positive
                            # Calculate RGB values
                            r = (color_int >> 16) & 0xFF
                            g = (color_int >> 8) & 0xFF
                            b = color_int & 0xFF
                            run.font.color.rgb = RGBColor(r, g, b)

                    elif child.tag == 'space':
                        # Add a space
                        run = paragraph.add_run(" ")
                    elif child.tag == 'image':
                        # Add an image
                        image_data = child.get('imageData')
                        if image_data:
                            image_bytes = base64.b64decode(image_data)
                            image_stream = io.BytesIO(image_bytes)
                            run = paragraph.add_run()
                            run.add_picture(image_stream)
            elif elem.tag == 'table':
                # Create the table
                column_count = int(elem.get('columnCount', '1'))
                rows = elem.findall('row')
                table = document.add_table(rows=len(rows), cols=column_count)
                for row_idx, row in enumerate(rows):
                    cells = row.findall('cell')
                    for col_idx, cell in enumerate(cells):
                        cell_paragraph = table.rows[row_idx].cells[col_idx].paragraphs[0]
                        paragraphs = cell.findall('paragraph')
                        for para in paragraphs:
                            # Process the paragraph content
                            for child in para:
                                if child.tag == 'content':
                                    # Get and format the text
                                    start_offset = int(child.get('startOffset', '0'))
                                    length = int(child.get('length', '0'))
                                    text = content_text[start_offset:start_offset+length]

                                    run = cell_paragraph.add_run(text)

                                    # Set the font size
                                    size = int(child.get('size', '12'))
                                    run.font.size = Pt(size)

                                    # Set bold and italic options
                                    if child.get('bold', 'false') == 'true':
                                        run.bold = True
                                    if child.get('italic', 'false') == 'true':
                                        run.italic = True

                                    # Set the font family
                                    font_family = child.get('family')
                                    if font_family:
                                        run.font.name = font_family

                                    # Set the color
                                    foreground = child.get('foreground')
                                    if foreground:
                                        color_int = int(foreground)
                                        if color_int < 0:
                                            color_int += 2**32  # Convert negative values to positive
                                        # Calculate RGB values
                                        r = (color_int >> 16) & 0xFF
                                        g = (color_int >> 8) & 0xFF
                                        b = color_int & 0xFF
                                        run.font.color.rgb = RGBColor(r, g, b)

                                elif child.tag == 'space':
                                    # Add a space
                                    run = cell_paragraph.add_run(" ")
                                elif child.tag == 'image':
                                    # Add an image
                                    image_data = child.get('imageData')
                                    if image_data:
                                        image_bytes = base64.b64decode(image_data)
                                        image_stream = io.BytesIO(image_bytes)
                                        run = cell_paragraph.add_run()
                                        run.add_picture(image_stream)
        else:
            print("'elements' could not be found in the XML.")

    # Save the document
    document.save(docx_file)

def main():
    if len(sys.argv) < 2:
        print("Usage: python udf_to_docx.py input.udf")
        exit()

    udf_file = sys.argv[1]

    if not os.path.isfile(udf_file):
        print(f"Input file not found: {udf_file}")
        exit()

    filename, ext = os.path.splitext(udf_file)

    if ext.lower() == '.udf':
        docx_file = filename + '.docx'
        udf_to_docx(udf_file, docx_file)
        print(f"DOCX file created: {docx_file}")
    else:
        print("Please provide a .udf file.")

if __name__ == '__main__':
    main()
